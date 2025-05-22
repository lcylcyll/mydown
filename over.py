# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Cm, Emu, Pt, Inches # Inches might be useful for conversion factor
# Make sure WD_SECTION is imported
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
import os
import copy # Not strictly needed for regenerate merge, but keep if used elsewhere
from lxml import etree
from PIL import Image # Requires Pillow library: pip install Pillow

# Requires python-docx: pip install python-docx
# Requires lxml: pip install lxml
# docxcompose is NOT used
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import traceback # For better error reporting

# ======================== 配置常量 ========================
A4_WIDTH_CM = 21
A4_HEIGHT_CM = 29.7
IMAGE_FOLDER = "processed"
OUTPUT_FOLDER = "output_docs"
MARGIN_CM = 0.1
PAGE_HEIGHT_BUFFER_PT = 13 # Using the value that worked for single files
PAGE_HEIGHT_BUFFER_EMU = Pt(PAGE_HEIGHT_BUFFER_PT).emu
A4_WIDTH = Cm(A4_WIDTH_CM)
A4_HEIGHT = Cm(A4_HEIGHT_CM)
MARGIN = Cm(MARGIN_CM)

# Conversion factor: EMU per Centimeter (approx)
EMU_PER_CM = 914400 / 2.54

# ======================== 工具函数 ========================

def emu_to_cm_display(emu_val):
    """辅助函数：将EMU整数转换为用于显示的厘米字符串"""
    if isinstance(emu_val, int) and emu_val >= 0:
        if EMU_PER_CM != 0: return f"{emu_val / EMU_PER_CM:.2f}cm"
        else: return f"InvalidEMUFactor({emu_val})"
    return f"InvalidEMU({emu_val})"

def set_cell_margins(cell, **kwargs):
    """设置表格单元格的边距 (pt)"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None: tcMar = OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for key, value in kwargs.items():
        if key in ["top", "bottom", "left", "right"]:
            margin_tag_str = f'w:{key}'; margin_qn_tag = qn(margin_tag_str)
            margin_el = tcMar.find(margin_qn_tag)
            if value is None or value <= 0:
                 if margin_el is not None and margin_el.getparent() == tcMar: tcMar.remove(margin_el)
            else:
                 if margin_el is None: margin_el = OxmlElement(margin_tag_str); tcMar.append(margin_el)
                 margin_value_twips = str(int(Pt(value).twips))
                 margin_el.set(qn('w:w'), margin_value_twips); margin_el.set(qn('w:type'), 'dxa')

def get_image_size_pixels(img_path):
    """获取图片的像素尺寸"""
    try:
        with Image.open(img_path) as img: return img.size
    except FileNotFoundError: raise FileNotFoundError(f"图片文件不存在: {img_path}")
    except Exception as e: raise RuntimeError(f"无法读取图片尺寸 {os.path.basename(img_path)}: {e}")

# --- Using the set_image_rotation function from your "working visual" version ---
def set_image_rotation(run, degrees):
    """设置图片旋转角度（仅支持90°倍数） - Simplified structure traversal"""
    # This function now receives the specific degree for the current image
    if not isinstance(degrees, (int, float)) or degrees % 90 != 0:
        print(f"警告: 无效或非90倍数旋转角度 {degrees}° 提供给 set_image_rotation。")
        return
    degrees = int(degrees) % 360
    if degrees == 0: return # No rotation needed (e.g., if passed 360)
    try:
        r = run._r
        drawing = r.find(qn('w:drawing'))
        if drawing is None: print("警告: run 中未找到 w:drawing 元素，无法旋转。"); return
        pic = None
        inline = drawing.find(qn('wp:inline'))
        if inline is not None:
            graphic = inline.find(qn('a:graphic'))
            if graphic is not None:
                graphicData = graphic.find(qn('a:graphicData'))
                if graphicData is not None: pic = graphicData.find(qn('pic:pic'))
        if pic is None:
            anchor = drawing.find(qn('wp:anchor'))
            if anchor is not None:
                 graphic = anchor.find(qn('a:graphic'))
                 if graphic is not None:
                     graphicData = graphic.find(qn('a:graphicData'))
                     if graphicData is not None: pic = graphicData.find(qn('pic:pic'))
        if pic is None: print(f"警告: drawing 内未找到 pic:pic 元素，无法旋转。"); return
        spPr = pic.find(qn('pic:spPr'))
        if spPr is None:
             spPr = OxmlElement('pic:spPr')
             blipFill = pic.find(qn('pic:blipFill'))
             nvPicPr = pic.find(qn('pic:nvPicPr'))
             if blipFill is not None: blipFill.addnext(spPr)
             elif nvPicPr is not None: nvPicPr.addnext(spPr)
             else: pic.insert(0, spPr)
        xfrm = spPr.find(qn('a:xfrm'))
        if xfrm is None: xfrm = OxmlElement('a:xfrm'); spPr.append(xfrm)
        rot_val = str(degrees * 60000)
        xfrm.set('rot', rot_val)
        print(f"      尝试设置具体旋转角度为 {degrees}° (rot={rot_val})") # More specific log
    except Exception as e: print(f"*** 设置图片旋转时出错: {e}")

def validate_image_path(filename):
    """验证并构建图片路径"""
    base_name = os.path.basename(str(filename))
    if not base_name: raise ValueError("提供的文件名无效或为空。")
    full_path = os.path.join(IMAGE_FOLDER, base_name)
    if not os.path.exists(full_path): raise FileNotFoundError(f"图片文件不存在: {full_path}")
    if not os.path.isfile(full_path): raise FileNotFoundError(f"路径不是一个文件: {full_path}")
    return full_path

# ======================== 布局生成函数 ========================
# --- Updated layout functions to handle rotate being None, int, or list ---

def create_full_page(doc, img_paths, rotate=None): # rotate can be None, int, or list (list len=1)
    """全页布局 - 拉伸填充 (Handles per-image rotation via list)"""
    if not img_paths: raise ValueError("full_page 布局需要图片文件。")

    section = doc.sections[-1]
    page_w = section.page_width; page_h = section.page_height
    margin_l = section.left_margin; margin_r = section.right_margin
    margin_t = section.top_margin; margin_b = section.bottom_margin

    available_width = max(Emu(1), page_w - margin_l - margin_r)
    available_height = max(Emu(1), page_h - margin_t - margin_b)

    if available_width <= Emu(1) or available_height <= Emu(1):
        raise ValueError(f"可用页面尺寸无效或过小 (W:{available_width}, H:{available_height})。")

    img_path = validate_image_path(img_paths[0])
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0)
    run = paragraph.add_run()
    pic = None

    # --- Determine Rotation for THIS image (index 0) ---
    current_rotate = None
    if isinstance(rotate, list):
        if len(rotate) > 0: current_rotate = rotate[0]
        else: print("    警告: full_page 提供了空的旋转列表。")
    elif isinstance(rotate, (int, float)):
        current_rotate = rotate # Apply single value
    print(f"  Full Page ({img_paths[0]}): Applying Rotation Logic for: {current_rotate}")
    # -----------------------------------------

    add_width = available_width
    add_height = available_height

    # <<< Apply dimension swap based on CURRENT_ROTATE >>>
    is_rotated_90_270 = current_rotate and isinstance(current_rotate, (int, float)) and current_rotate % 180 != 0
    if is_rotated_90_270:
        print(f"    Swapping dimensions for add_picture (Rotation: {current_rotate}°).")
        add_width = available_height
        add_height = available_width

    print(f"    Adding picture with initial W={emu_to_cm_display(add_width)}, H={emu_to_cm_display(add_height)}")
    try:
        pic = run.add_picture(img_path, width=add_width, height=add_height)
    except Exception as e:
        print(f"*** Error adding picture in create_full_page: {e}")
        run.add_text(f"[Error adding {img_paths[0]}]")
        return

    # <<< Attempt rotation based on CURRENT_ROTATE >>>
    if current_rotate and pic:
        try: set_image_rotation(run, current_rotate)
        except Exception as e: print(f"*** Error applying rotation in create_full_page: {e}")

def create_vertical_split(doc, img_paths, split_num, rotate=None): # rotate can be None, int, or list
    """垂直分栏布局 (表格) - Handles per-image rotation via list"""
    if not isinstance(split_num, int) or split_num <= 0: raise ValueError("split_num must be positive integer")
    original_filenames = list(img_paths) # Keep original list for reference
    files_to_process = list(img_paths) # Use a copy for processing/trimming
    if len(files_to_process) > split_num: print(f"警告: V-Split 图片数 ({len(files_to_process)}) > 分栏数 ({split_num}). Trimming."); files_to_process = files_to_process[:split_num]
    elif len(files_to_process) < split_num: print(f"警告: V-Split 图片数 ({len(files_to_process)}) < 分栏数 ({split_num}). Empty cells.")

    section = doc.sections[-1]
    available_width = max(Emu(1), section.page_width - section.left_margin - section.right_margin)
    _available_height = max(Emu(1), section.page_height - section.top_margin - section.bottom_margin)
    if available_width <= Emu(1) or _available_height <= Emu(1): raise ValueError("Available page size non-positive.")

    available_height = max(Emu(1), _available_height - PAGE_HEIGHT_BUFFER_EMU)
    if available_height <= Emu(1): print(f"警告: 应用缓冲区 ({PAGE_HEIGHT_BUFFER_PT}pt) 后垂直可用高度 ({available_height} EMU) 无效。使用原始高度。"); available_height = _available_height

    row_height_emu = max(Emu(1), int(available_height / split_num))
    print(f"  Vertical Split ({split_num}): Row H={emu_to_cm_display(row_height_emu)} (Buffer={PAGE_HEIGHT_BUFFER_PT}pt)")

    table = doc.add_table(rows=split_num, cols=1)
    table.autofit = False; table.allow_autofit = False
    table.width = available_width

    for i in range(split_num): # Iterate through cells
        cell = table.cell(i, 0); cell.width = available_width
        row = table.rows[i]; row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        row.height = row_height_emu
        set_cell_margins(cell, top=Pt(0), bottom=Pt(0), left=Pt(0), right=Pt(0))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        while len(cell.paragraphs) > 1: cell._element.remove(cell.paragraphs[-1]._p)
        paragraph = cell.paragraphs[0]; paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = paragraph.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0

        if i < len(files_to_process): # Check if there's an image for this cell index
            img_file = files_to_process[i]
            run = paragraph.add_run(); pic = None
            try:
                img_path = validate_image_path(img_file)
                target_cell_width = available_width
                target_cell_height = row_height_emu

                # --- Determine Rotation for THIS image ---
                current_rotate = None
                if isinstance(rotate, list):
                    # Use original_filenames index which matches config list index
                    original_index = -1
                    try:
                       # Find the index of the current img_file in the original list
                       original_index = original_filenames.index(img_file) # Assumes unique filenames in config!
                       if original_index < len(rotate):
                           current_rotate = rotate[original_index]
                       else: print(f"    警告: 旋转列表长度不足 (图: {img_file}, 原索引: {original_index})。")
                    except ValueError:
                       print(f"    警告: 无法在原始文件列表找到 {img_file} 以获取旋转索引。") # Should not happen if files_to_process comes from original
                    except IndexError: # Should be caught by length check above, but belt-and-suspenders
                       print(f"    警告: 计算得到的旋转索引 {original_index} 超出范围。")
                elif isinstance(rotate, (int, float)): # Apply single value to all
                    current_rotate = rotate
                # If rotate is None initially, current_rotate remains None
                print(f"    V-Split Cell {i} ({img_file}): Applying Rotation Logic for: {current_rotate}")
                # -----------------------------------------

                add_width = target_cell_width
                add_height = target_cell_height

                # <<< Apply dimension swap based on CURRENT_ROTATE >>>
                is_rotated_90_270 = current_rotate and isinstance(current_rotate, (int, float)) and current_rotate % 180 != 0
                if is_rotated_90_270:
                    print(f"      Swapping dimensions for add_picture (Rotation: {current_rotate}°).")
                    add_width = target_cell_height
                    add_height = target_cell_width

                print(f"      Adding picture with initial W={emu_to_cm_display(add_width)}, H={emu_to_cm_display(add_height)}")
                pic = run.add_picture(img_path, width=add_width, height=add_height)

                # <<< Attempt rotation based on CURRENT_ROTATE >>>
                if current_rotate and pic:
                     set_image_rotation(run, current_rotate)

            except (FileNotFoundError, RuntimeError, ValueError) as e:
                paragraph.clear(); paragraph.add_run(f"[Error loading {img_file}: {e}]")
                print(f"*** Error processing image {img_file} in V-split: {e}")
            except Exception as e:
                 paragraph.clear(); paragraph.add_run(f"[Error processing {img_file}]")
                 print(f"*** Unexpected Error processing image {img_file} in V-split: {e}")
        else:
            paragraph.add_run("") # Keep empty cells as empty paragraphs

def create_custom_ratio(doc, img_paths, ratios, orientation, rotate=None): # rotate can be None, int, or list
    """自定义比例布局 (表格) - Handles per-image rotation via list"""
    original_filenames = list(img_paths) # Keep original list for reference
    files_to_process = list(img_paths) # Use a copy for processing/trimming (though less likely needed here)

    section = doc.sections[-1]
    available_width = max(Emu(1), section.page_width - section.left_margin - section.right_margin)
    _available_height = max(Emu(1), section.page_height - section.top_margin - section.bottom_margin)
    if available_width <= Emu(1) or _available_height <= Emu(1): raise ValueError("Available page size non-positive.")

    if orientation == "vertical":
        available_height = max(Emu(1), _available_height - PAGE_HEIGHT_BUFFER_EMU)
        if available_height <= Emu(1): print(f"警告: 应用缓冲区后垂直可用高度 ({available_height} EMU) 无效。使用原始高度。"); available_height = _available_height
    else: # Horizontal
        available_height = _available_height

    num_items = len(files_to_process) # Use length of files actually being processed
    if not ratios or len(ratios) != num_items: raise ValueError(f"Ratios list length ({len(ratios) if ratios else 0}) doesn't match image count ({num_items}).")

    try:
        ratios_float = [float(r) for r in ratios]; total_ratio = sum(ratios_float)
        if total_ratio <= 0: raise ValueError("Ratios sum must be positive.")
        normalized_ratios = [r / total_ratio for r in ratios_float]
    except (TypeError, ValueError): raise ValueError("Ratios must be numeric.")

    if orientation == "horizontal":
        print(f"  Custom Ratio (Horizontal): Ratios={ratios}, Avail W={emu_to_cm_display(available_width)}, Avail H={emu_to_cm_display(available_height)}")
        table = doc.add_table(rows=1, cols=num_items)
        table.autofit = False; table.allow_autofit = False; table.width = available_width
        col_widths_emu = [max(Emu(1), int(available_width * ratio)) for ratio in normalized_ratios]
        col_widths_emu[-1] = max(Emu(1), available_width - sum(col_widths_emu[:-1]))
        for i, width in enumerate(col_widths_emu): table.columns[i].width = width; print(f"    H-Ratio Col {i}: Width={emu_to_cm_display(width)}")
        table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[0].height = available_height

        for i, img_file in enumerate(files_to_process): # Iterate through images to process
            cell = table.cell(0, i)
            cell_width = col_widths_emu[i]; cell_height = available_height
            set_cell_margins(cell, top=Pt(0), bottom=Pt(0), left=Pt(0), right=Pt(0))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            while len(cell.paragraphs) > 1: cell._element.remove(cell.paragraphs[-1]._p)
            paragraph = cell.paragraphs[0]; paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = paragraph.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
            run = paragraph.add_run(); pic = None
            try:
                img_path = validate_image_path(img_file)

                # --- Determine Rotation for THIS image ---
                current_rotate = None
                if isinstance(rotate, list):
                    original_index = -1
                    try: original_index = original_filenames.index(img_file); current_rotate = rotate[original_index] if original_index < len(rotate) else None
                    except (ValueError, IndexError): print(f"    警告: H-Ratio 无法确定图像 {img_file} 的旋转。")
                elif isinstance(rotate, (int, float)): current_rotate = rotate
                print(f"    H-Ratio Cell {i} ({img_file}): Applying Rotation Logic for: {current_rotate}")
                # -----------------------------------------

                add_width = cell_width
                add_height = cell_height

                # <<< Apply dimension swap based on CURRENT_ROTATE >>>
                is_rotated_90_270 = current_rotate and isinstance(current_rotate, (int, float)) and current_rotate % 180 != 0
                if is_rotated_90_270:
                    print(f"      Swapping dimensions for add_picture (Rotation: {current_rotate}°).")
                    add_width = cell_height
                    add_height = cell_width

                print(f"      Adding picture with initial W={emu_to_cm_display(add_width)}, H={emu_to_cm_display(add_height)}")
                pic = run.add_picture(img_path, width=add_width, height=add_height)

                # <<< Attempt rotation based on CURRENT_ROTATE >>>
                if current_rotate and pic: set_image_rotation(run, current_rotate)

            except (FileNotFoundError, RuntimeError, ValueError) as e: paragraph.clear(); paragraph.add_run(f"[Error: {e}]"); print(f"*** Error processing image {img_file} in H-ratio: {e}")
            except Exception as e: paragraph.clear(); paragraph.add_run(f"[Error processing {img_file}]"); print(f"*** Unexpected Error processing image {img_file} in H-ratio: {e}")

    elif orientation == "vertical":
        print(f"  Custom Ratio (Vertical): Ratios={ratios}. Avail W={emu_to_cm_display(available_width)}, Buffered Avail H={emu_to_cm_display(available_height)}")
        table = doc.add_table(rows=num_items, cols=1)
        table.autofit = False; table.allow_autofit = False; table.width = available_width
        row_heights_emu = [max(Emu(1), int(available_height * ratio)) for ratio in normalized_ratios]
        row_heights_emu[-1] = max(Emu(1), available_height - sum(row_heights_emu[:-1]))

        for i, img_file in enumerate(files_to_process): # Iterate through images to process
            cell = table.cell(i, 0); cell.width = available_width
            row = table.rows[i]; row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            current_row_height = row_heights_emu[i]
            row.height = current_row_height; print(f"    V-Ratio Row {i}: Height={emu_to_cm_display(current_row_height)}")
            set_cell_margins(cell, top=Pt(0), bottom=Pt(0), left=Pt(0), right=Pt(0))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            while len(cell.paragraphs) > 1: cell._element.remove(cell.paragraphs[-1]._p)
            paragraph = cell.paragraphs[0]; paragraph.clear()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = paragraph.paragraph_format; pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
            run = paragraph.add_run(); pic = None
            try:
                img_path = validate_image_path(img_file)
                cell_width = available_width; cell_height = current_row_height

                # --- Determine Rotation for THIS image ---
                current_rotate = None
                if isinstance(rotate, list):
                    original_index = -1
                    try: original_index = original_filenames.index(img_file); current_rotate = rotate[original_index] if original_index < len(rotate) else None
                    except (ValueError, IndexError): print(f"    警告: V-Ratio 无法确定图像 {img_file} 的旋转。")
                elif isinstance(rotate, (int, float)): current_rotate = rotate
                print(f"    V-Ratio Cell {i} ({img_file}): Applying Rotation Logic for: {current_rotate}")
                # -----------------------------------------

                add_width = cell_width
                add_height = cell_height

                # <<< Apply dimension swap based on CURRENT_ROTATE >>>
                is_rotated_90_270 = current_rotate and isinstance(current_rotate, (int, float)) and current_rotate % 180 != 0
                if is_rotated_90_270:
                    print(f"      Swapping dimensions for add_picture (Rotation: {current_rotate}°).")
                    add_width = cell_height
                    add_height = cell_width

                print(f"      Adding picture with initial W={emu_to_cm_display(add_width)}, H={emu_to_cm_display(add_height)}")
                pic = run.add_picture(img_path, width=add_width, height=add_height)

                # <<< Attempt rotation based on CURRENT_ROTATE >>>
                if current_rotate and pic: set_image_rotation(run, current_rotate)

            except (FileNotFoundError, RuntimeError, ValueError) as e: paragraph.clear(); paragraph.add_run(f"[Error: {e}]"); print(f"*** Error processing image {img_file} in V-ratio: {e}")
            except Exception as e: paragraph.clear(); paragraph.add_run(f"[Error processing {img_file}]"); print(f"*** Unexpected Error processing image {img_file} in V-ratio: {e}")
    else: raise ValueError(f"Invalid orientation for custom_ratio: {orientation}")


# ======================== 主生成函数 ========================
def generate_single_doc(docx_path, config):
    """生成单个文档 (saves to docx_path) - Uses updated layout functions"""
    doc = Document()
    section = doc.sections[0]

    page_orientation_setting = config.get("orientation", "vertical").lower()
    if page_orientation_setting == "horizontal":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = A4_HEIGHT; section.page_height = A4_WIDTH
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = A4_WIDTH; section.page_height = A4_HEIGHT

    section.left_margin = MARGIN; section.right_margin = MARGIN
    section.top_margin = MARGIN; section.bottom_margin = MARGIN

    img_files_basenames = config.get("files", [])
    if not img_files_basenames: raise ValueError("配置中 'files' 列表为空")

    try: validate_image_path(img_files_basenames[0])
    except (FileNotFoundError, ValueError) as e: raise RuntimeError(f"主要图片文件验证失败 ({img_files_basenames[0]}) for {os.path.basename(docx_path)}: {str(e)}")

    layout = config.get("layout", "full_page")
    # Pass the rotate setting (None, int, or list) from preprocessed config
    rotate = config.get("rotate")

    try:
        print(f"\nGenerating {os.path.basename(docx_path)} (Layout: '{layout}', Page: {page_orientation_setting}, Rotate: {rotate})...")
        if layout == "full_page":
            create_full_page(doc, img_files_basenames, rotate) # Pass rotate setting
        elif layout.startswith("vertical_") and layout.endswith("_split"):
            try:
                 split_num = int(layout.split("_")[1])
                 if split_num <= 0 : raise ValueError("Split number must be positive")
            except (IndexError, ValueError, TypeError):
                 raise ValueError(f"垂直分割格式无效: '{layout}' (应为 vertical_N_split, N > 0)")
            create_vertical_split(doc, img_files_basenames, split_num, rotate) # Pass rotate setting
        elif layout == "custom_ratio":
            split_orientation = config.get("ratio_orientation", page_orientation_setting).lower()
            ratios = config.get("ratios")
            if not ratios: raise ValueError("custom_ratio requires 'ratios' list.")
            create_custom_ratio(doc, img_files_basenames, ratios, split_orientation, rotate) # Pass rotate setting
        else: raise ValueError(f"Unknown layout type: '{layout}'")

    except Exception as e:
        print(f"\n*** Layout generation failed for {os.path.basename(docx_path)} ({layout}) ***")
        traceback.print_exc()
        raise RuntimeError(f"布局生成失败 for {os.path.basename(docx_path)} ({layout}): {str(e)}")

    try:
        os.makedirs(os.path.dirname(docx_path), exist_ok=True)
        doc.save(docx_path)
    except Exception as e: raise RuntimeError(f"无法保存文档 {docx_path}: {e}")


# ======================== 文档合并 (Regenerate Content) V4 ========================
# --- USING THE REGENERATE MERGE FUNCTION ---
def merge_documents_regenerate(input_paths_in_order, output_path, config_map):
    """
    Merges documents by regenerating the content for each page directly
    into the master document using the original layout functions.
    Uses SECTION breaks.
    """
    if not input_paths_in_order: print("没有提供用于合并的输入文件路径。"); return False

    master = Document()
    config_lookup = {os.path.basename(k): v for k, v in config_map.items()}
    first_doc_key = os.path.basename(input_paths_in_order[0])
    first_config = config_lookup.get(first_doc_key)
    master_sec = master.sections[0]

    if first_config:
        print(f"\n合并 V4 (Regenerate): 使用 '{first_doc_key}' 的配置设置页面基础。")
        page_orientation_setting = first_config.get("orientation", "vertical").lower()
        if page_orientation_setting == "horizontal":
            master_sec.orientation = WD_ORIENT.LANDSCAPE
            master_sec.page_width = A4_HEIGHT; master_sec.page_height = A4_WIDTH
        else:
            master_sec.orientation = WD_ORIENT.PORTRAIT
            master_sec.page_width = A4_WIDTH; master_sec.page_height = A4_HEIGHT
        master_sec.left_margin = MARGIN; master_sec.right_margin = MARGIN
        master_sec.top_margin = MARGIN; master_sec.bottom_margin = MARGIN
    else: print(f"警告: 无法找到第一个文档 '{first_doc_key}' 的配置。将使用默认页面设置。")

    if len(master.element.body) > 0 and master.element.body[0].tag == qn('w:p') and not master.element.body[0].xpath('.//w:r | .//w:drawing | .//w:t'):
         print("  移除初始空白段落。")
         master.element.body.remove(master.element.body[0])

    print(f"\n开始按顺序重新生成内容到主文档...")
    merge_successful = True

    for i, path in enumerate(input_paths_in_order):
        current_filename_key = os.path.basename(path)
        config = config_lookup.get(current_filename_key)

        if not config:
            print(f"*** 错误: 未找到文档 '{current_filename_key}' 的配置信息，无法重新生成内容。跳过。")
            try: master.add_paragraph(f"[错误: 未找到 {current_filename_key} 的配置，内容缺失]")
            except Exception as pe: print(f"*** 无法添加错误段落: {pe}")
            merge_successful = False
            if i < len(input_paths_in_order) - 1:
                 try: master.add_section(WD_SECTION.NEW_PAGE)
                 except Exception as se: print(f"*** 无法添加分节符: {se}")
            continue

        print(f"  ({i+1}/{len(input_paths_in_order)}) 正在重新生成内容来源: {current_filename_key}")
        layout = config.get("layout", "full_page")
        img_files_basenames = config.get("files", [])
        rotate = config.get("rotate") # Pass the potentially list-based rotate setting
        ratios = config.get("ratios")
        ratio_orientation = config.get("ratio_orientation", config.get("orientation", "vertical")).lower()

        try:
            # Call layout functions which now handle per-image rotation
            if layout == "full_page": create_full_page(master, img_files_basenames, rotate)
            elif layout.startswith("vertical_") and layout.endswith("_split"):
                try: split_num = int(layout.split("_")[1]); assert split_num > 0
                except: raise ValueError(f"合并期间发现垂直分割格式无效: '{layout}'")
                create_vertical_split(master, img_files_basenames, split_num, rotate)
            elif layout == "custom_ratio":
                if not ratios: raise ValueError("合并期间发现 custom_ratio 缺少 'ratios'")
                create_custom_ratio(master, img_files_basenames, ratios, ratio_orientation, rotate)
            else: raise ValueError(f"合并期间发现未知布局类型: '{layout}'")
            print(f"    '{current_filename_key}' 内容生成成功。")

        except Exception as e:
            print(f"*** 错误: 重新生成 '{current_filename_key}' 内容时失败: {e}")
            traceback.print_exc(); merge_successful = False
            try: master.add_paragraph(f"[错误: 重新生成 {current_filename_key} 失败: {e}]")
            except Exception as pe: print(f"*** 无法添加错误段落: {pe}")

        if i < len(input_paths_in_order) - 1:
            try: print(f"    添加分节符 (下一页)。"); master.add_section(WD_SECTION.NEW_PAGE)
            except Exception as se: print(f"*** 添加分节符时出错: {se}"); merge_successful = False

    try:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        master.save(output_path)
        if merge_successful: print(f"\n合并 V4 (Regenerate) 完成! 输出文件: {output_path}")
        else: print(f"\n*** 合并 V4 (Regenerate) 过程中出现错误，输出文件可能不完整或损坏: {output_path}")
        return merge_successful
    except Exception as e: print(f"*** 保存最终合并文档时出错 {output_path}: {e}"); traceback.print_exc(); return False


# ======================== 配置预处理 ========================
# --- UPDATED preprocess_config to handle rotate list ---
def preprocess_config(raw_config):
    """清洗配置数据, 使用 basename for files. Handles list or single value for rotate."""
    print("开始详细配置预处理...")
    processed_config = {}
    processed_count = 0
    skipped_count = 0
    if not isinstance(raw_config, dict): raise TypeError("raw_config must be a dictionary.")

    for filename, cfg in raw_config.items():
        print(f"  处理配置: '{filename}'")
        is_valid = True
        if not isinstance(cfg, dict): print(f"    错误: 配置条目 '{filename}' 不是字典。跳过。"); skipped_count += 1; continue
        if "files" not in cfg or not isinstance(cfg["files"], list): print(f"    错误: '{filename}' 配置缺少 'files' 列表或类型错误。跳过。"); is_valid = False
        if "layout" not in cfg: print(f"    错误: '{filename}' 配置缺少 'layout'。跳过。"); is_valid = False
        if not is_valid: skipped_count += 1; continue

        clean_files = []
        invalid_files = []
        for f in cfg.get("files", []):
             if f and isinstance(f, str) and str(f).strip(): clean_files.append(os.path.basename(str(f).strip()))
             else: invalid_files.append(f)
        if invalid_files: print(f"    警告: '{filename}' 中的文件条目无效/为空: {invalid_files}。")
        if not clean_files: print(f"    错误: '{filename}' 清理后 'files' 列表为空。跳过。"); skipped_count += 1; continue

        page_orientation = cfg.get("orientation", "vertical").lower()
        if page_orientation not in ["vertical", "horizontal"]: page_orientation = "vertical"
        print(f"    页面方向: {page_orientation}")
        ratio_orientation = cfg.get("ratio_orientation", page_orientation).lower()
        if ratio_orientation not in ["vertical", "horizontal"]: ratio_orientation = page_orientation
        print(f"    比例方向: {ratio_orientation}")

        rotate_config_val = cfg.get("rotate")
        print(f"    原始旋转值: {rotate_config_val}")
        final_rotate_setting = None

        if rotate_config_val is None: print(f"    无旋转值。"); final_rotate_setting = None
        elif isinstance(rotate_config_val, (int, float)):
            try:
                rotate_int = int(rotate_config_val)
                if rotate_int % 90 == 0: final_rotate_setting = rotate_int % 360; print(f"    有效单旋转值: {final_rotate_setting}°")
                else: print(f"    警告: 单旋转值 '{rotate_config_val}' 非90倍数。忽略旋转。")
            except: print(f"    警告: 单旋转值 '{rotate_config_val}' 无效。忽略旋转。")
        elif isinstance(rotate_config_val, list):
            if len(rotate_config_val) == len(clean_files):
                validated_rotations = []
                valid_list = True
                for idx, r_val in enumerate(rotate_config_val):
                    if r_val is None: validated_rotations.append(None)
                    elif isinstance(r_val, (int, float)):
                        try:
                            r_int = int(r_val)
                            if r_int % 90 == 0: validated_rotations.append(r_int % 360)
                            else: print(f"    警告: 列表旋转值 '{r_val}' (索引 {idx}) 非90倍数。设为 None。"); validated_rotations.append(None); valid_list = False
                        except: print(f"    警告: 列表旋转值 '{r_val}' (索引 {idx}) 无效。设为 None。"); validated_rotations.append(None); valid_list = False
                    else: print(f"    警告: 列表旋转值 '{r_val}' (索引 {idx}) 类型错误。设为 None。"); validated_rotations.append(None); valid_list = False
                final_rotate_setting = validated_rotations
                print(f"    有效列表旋转值: {final_rotate_setting}")
                # if not valid_list: print(f"    注意: 旋转列表包含无效值，已替换为 None。") # Optional warning
            else: print(f"    错误: 'rotate' 列表长度 ({len(rotate_config_val)}) 与 'files' 列表长度 ({len(clean_files)}) 不匹配。忽略旋转。"); final_rotate_setting = None
        else: print(f"    警告: 'rotate' 值类型未知 ({type(rotate_config_val).__name__})。忽略旋转。")

        ratios_list = cfg.get("ratios"); clean_ratios = []
        print(f"    原始比例列表: {ratios_list}")
        if isinstance(ratios_list, list):
             valid_ratios = True
             for r in ratios_list:
                 try: float_r = float(r); assert float_r > 0; clean_ratios.append(float_r)
                 except: print(f"    错误: '{filename}' 比例值 '{r}' 无效。比例列表作废。"); valid_ratios = False; break
             if not valid_ratios: clean_ratios = []
             else: print(f"    有效比例列表: {clean_ratios}")
        else:
             if ratios_list is not None: print(f"    警告: '{filename}' 的 'ratios' 不是列表。忽略。")

        processed_config[filename] = {
            "files": clean_files, "layout": cfg["layout"], "ratios": clean_ratios,
            "orientation": page_orientation, "ratio_orientation": ratio_orientation,
            "rotate": final_rotate_setting # Store None, int, or list
        }
        processed_count += 1
        print(f"  配置 '{filename}' 处理完成。")

    print(f"配置预处理结束。有效处理: {processed_count}，跳过: {skipped_count}。")
    return processed_config


# ======================== 主程序 ========================
if __name__ == "__main__":
    # --- 1. Setup ---
    print("脚本开始执行...")
    if not os.path.isdir(IMAGE_FOLDER): print(f"错误: 图片文件夹 '{IMAGE_FOLDER}' 不存在。"); exit(1)
    try: os.makedirs(OUTPUT_FOLDER, exist_ok=True); print(f"输出文件夹: '{os.path.abspath(OUTPUT_FOLDER)}'")
    except OSError as e: print(f"错误: 无法创建输出文件夹 '{OUTPUT_FOLDER}': {e}"); exit(1)

    # --- 2. Configuration ---
    RAW_CONFIG = {
         "01.docx": {"files": ["01.tif"], "layout": "full_page", "rotate": 90},
        "03.docx": {"files": ["1.tif", "2.tif", "3.tif"], "layout": "vertical_3_split"},
        #"02.docx": { "files": ["02.tif", "002.tif"], "layout": "custom_ratio", "ratios": [0.8, 0.2], "orientation": "vertical", "ratio_orientation": "horizontal"},
        "02.docx": { "files": ["02.tif", "002.tif"], "layout": "custom_ratio", "ratios": [0.8, 0.2], "orientation": "vertical", "ratio_orientation": "vertical"},
        "2.docx": {"files": ["4.tif", "5.tif", "6.tif"], "layout": "vertical_3_split"},
        "3.docx": {"files": ["7.tif", "8.tif", "9.tif"], "layout": "vertical_3_split"},
        "4.docx": {"files": ["10.tif", "11.tif"], "layout": "vertical_2_split", "rotate": 90},
        "5.docx": {"files": ["12.tif", "013.tif", "13.tif"], "layout": "vertical_3_split"},
        "6.docx": {"files": ["014.tif", "14.tif", "15.tif"], "layout": "vertical_3_split"},
        "16.docx": {"files": ["16.tif"], "layout": "full_page"},
        "7.docx": {"files": ["17.tif", "p003.tif", "p040.tif"], "layout": "vertical_3_split"}, # Check p03.tif exists
        "8.docx": {"files": ["18.tif", "n290.tif"], "layout": "vertical_2_split", "rotate": [None, 90]}, # MODIFIED for per-image rotation
        "9.docx": {"files": ["v097.tif", "v098.tif"], "layout": "vertical_2_split"},
        "30.docx": {"files": ["30.tif"], "layout": "full_page"},
        "080a.docx": {"files": ["080a.tif"], "layout": "full_page"},
        "805.docx": {"files": ["805.tif"], "layout": "full_page"},
        "a245.docx": {"files": ["a245.tif"], "layout": "full_page"},
        "61.docx": {"files": ["n61.tif"], "layout": "full_page"},
        "n62.docx": {"files": ["n62.tif"], "layout": "full_page"},
    }

    expected_file_count = len(RAW_CONFIG)
    print(f"\n预期生成文件数量: {expected_file_count}")
    processed_files_paths = []
    final_merged_base_name = "final_merged.docx"
    final_output_path = os.path.join(OUTPUT_FOLDER, final_merged_base_name)
    LAYOUT_CONFIG = {}
    generation_successful_overall = True
    generated_files_count = 0
    failed_generation_files = []
    processed_config_count = 0

    # --- 3. Preprocess Config and Generate Individual Docs ---
    try:
        print("\n正在预处理配置...")
        LAYOUT_CONFIG = preprocess_config(RAW_CONFIG) # Uses updated preprocess function
        processed_config_count = len(LAYOUT_CONFIG)
        print(f"配置处理完成。有效配置: {processed_config_count}/{expected_file_count}")

        if processed_config_count > 0:
             print("\n开始生成单个文档...")
             sorted_config_keys = sorted(LAYOUT_CONFIG.keys())

             for filename_key in sorted_config_keys:
                 config = LAYOUT_CONFIG[filename_key]
                 target_docx_path = os.path.join(OUTPUT_FOLDER, filename_key)
                 try:
                     # Calls generate_single_doc, which calls updated layout functions
                     generate_single_doc(target_docx_path, config)
                     print(f"--- 成功生成: {os.path.basename(target_docx_path)}")
                     processed_files_paths.append(target_docx_path)
                     generated_files_count += 1
                 except (FileNotFoundError, ValueError, RuntimeError, TypeError) as e:
                     print(f"*** 生成 {filename_key} 失败: {str(e)}")
                     generation_successful_overall = False; failed_generation_files.append(filename_key)
                 except Exception as e:
                     print(f"*** 生成 {filename_key} 时发生意外错误: {str(e)}"); traceback.print_exc()
                     generation_successful_overall = False; failed_generation_files.append(filename_key)

             print(f"\n单个文档生成完成。成功: {generated_files_count} / {processed_config_count} 个有效配置。")
             if failed_generation_files: print(f"失败的文件: {', '.join(failed_generation_files)}")
        else: print("\n没有有效的配置，无法生成任何文档。")

    except Exception as e:
         print(f"\n*** 程序在配置处理或生成循环控制阶段发生意外错误: {str(e)}"); traceback.print_exc()
         generation_successful_overall = False

    # --- 4. Merge and Cleanup ---
    RUN_MERGE_AND_CLEANUP = True  # <<< ENABLE MERGE

    print(f"\n合并与清理步骤当前状态: {'启用 (使用 Regenerate)' if RUN_MERGE_AND_CLEANUP else '禁用'}")

    if RUN_MERGE_AND_CLEANUP:
        if generated_files_count > 0 :
            if not generation_successful_overall: print("\n警告：由于部分文件生成失败，合并结果可能不完整或包含错误信息。")

            print(f"\n准备合并 {len(processed_files_paths)} 个成功生成的文档 (使用 Regenerate)...")
            processed_files_paths.sort()
            print("将按以下顺序合并:")
            for p in processed_files_paths: print(f"  - {os.path.basename(p)}")

            try:
                # --- USE THE REGENERATE MERGE FUNCTION ---
                merge_success = merge_documents_regenerate(processed_files_paths, final_output_path, LAYOUT_CONFIG)

                if merge_success:
                     if os.path.exists(final_output_path):
                          print(f"\n合并 (Regenerate) 完成! 输出文件: {final_output_path}")
                          # --- Cleanup ---
                          print("\n准备清理临时生成的单个文档...")
                          cleaned_count = 0; errors_cleaning = False
                          for f_path in processed_files_paths:
                              try:
                                  if os.path.exists(f_path): os.remove(f_path); cleaned_count += 1
                                  else: print(f"  警告: 尝试删除时未找到文件 {os.path.basename(f_path)}。")
                              except Exception as e: print(f"*** 删除临时文件 {os.path.basename(f_path)} 失败: {str(e)} ***"); errors_cleaning = True
                          print(f"清理完成。成功删除了 {cleaned_count} / {len(processed_files_paths)} 个文件。")
                          if errors_cleaning: print("清理过程中出现错误。")
                     else: print(f"\n*** 合并 (Regenerate) 报告成功但未找到输出文件: {final_output_path}。跳过清理。")
                # else: Failure message printed inside merge function

            except Exception as e:
                print(f"\n*** 合并文档 (Regenerate) 时发生意外错误: {str(e)}"); traceback.print_exc()
                print(f"  临时生成的单个文档保留在: {OUTPUT_FOLDER}")

        elif processed_config_count > 0: print("\n所有有效配置的文档生成均失败，无法进行合并或清理。")
        elif processed_config_count == 0: pass
        else: print("\n由于未知原因，跳过合并和清理。")
    else: print(f"\n已跳过合并和清理步骤。生成的单个文档保留在: {OUTPUT_FOLDER}")

    print("\n脚本执行完毕。")